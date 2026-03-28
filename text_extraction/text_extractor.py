"""
Text Extraction Module for Resume Skill Recognition System.
Handles extraction of text from PDF, DOCX, and image files.
"""

import io
from pathlib import Path
from typing import Dict, Optional, Union, Any
import PyPDF2
import pdfplumber
from docx import Document
from PIL import Image
import pytesseract

from utils import get_logger, config


logger = get_logger(__name__)


class TextExtractor:
    """Extracts text from various document formats with intelligent fallback."""
    
    def __init__(self):
        """Initialize text extractor with configuration."""
        self.min_text_length = config.get('extraction.min_text_length', 100)
        self.ocr_fallback = config.get('extraction.ocr_fallback', True)
        self.supported_formats = config.get(
            'extraction.supported_formats',
            ['pdf', 'docx', 'doc', 'png', 'jpg', 'jpeg', 'bmp', 'tiff', 'tif', 'webp']
        )
        self.image_extensions = {'.png', '.jpg', '.jpeg', '.bmp', '.tiff', '.tif', '.webp'}
        self.builtin_supported_extensions = {'.pdf', '.docx', '.doc'} | self.image_extensions
        self.ocr_language = config.get('extraction.ocr_language', 'eng')
        self.pdf_method = config.get('extraction.pdf_extraction_method', 'pdfplumber')
        self.use_pymupdf = config.get('extraction.use_pymupdf', True)
        self.use_camelot = config.get('extraction.use_camelot', True)
        self.capture_layout = config.get('extraction.capture_layout', True)
        self.ocr_min_chars = config.get('extraction.ocr_min_chars', 60)
        self.ocr_page_limit = config.get('extraction.ocr_page_limit', 5)
        self.poppler_path = config.get('extraction.poppler_path', r'C:\poppler\Library\bin')
        self.tesseract_path = config.get('extraction.tesseract_path', r'C:\Program Files\Tesseract-OCR\tesseract.exe')
        
        logger.info("TextExtractor initialized with OCR fallback: %s", self.ocr_fallback)
    
    def extract(self, file_path: Union[str, Path]) -> Dict[str, any]:
        """
        Extract text from document file.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Dictionary containing:
                - text: Extracted text
                - method: Extraction method used
                - success: Whether extraction succeeded
                - error: Error message if failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error("File not found: %s", file_path)
            return {
                'text': '',
                'method': 'none',
                'success': False,
                'error': 'File not found'
            }
        
        ext = ''.join(ch for ch in file_path.suffix.lower().strip() if ch.isalnum() or ch == '.')
        config_supported = {
            f".{fmt.lower().lstrip('.')}" for fmt in self.supported_formats
        }
        normalized_supported = config_supported | self.builtin_supported_extensions
        
        try:
            if ext not in normalized_supported:
                logger.warning("Unsupported file format: %s", ext)
                return {
                    'text': '',
                    'method': 'none',
                    'success': False,
                    'error': f'Unsupported format: {ext}'
                }

            if ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif ext in self.image_extensions:
                return self._extract_from_image(file_path)
            else:
                logger.warning("Unsupported file format: %s", ext)
                return {
                    'text': '',
                    'method': 'none',
                    'success': False,
                    'error': f'Unsupported format: {ext}'
                }
        except Exception as e:
            logger.error("Error extracting text from %s: %s", file_path, str(e))
            return {
                'text': '',
                'method': 'error',
                'success': False,
                'error': str(e)
            }

    def _extract_from_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text directly from image files using OCR.

        Args:
            file_path: Path to image file

        Returns:
            Extraction result dictionary
        """
        logger.info("Extracting text from image: %s", file_path.name)

        try:
            if self.tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_path

            with Image.open(file_path) as image:
                text = pytesseract.image_to_string(image, lang=self.ocr_language)

            # For images, accept any non-empty OCR text; strict minimum can drop valid short resumes.
            success = len(text.strip()) > 0
            return {
                'text': text,
                'method': 'image_ocr',
                'success': success,
                'error': None if success else 'Insufficient text extracted from image',
                'ocr_performed': True
            }
        except Exception as e:  # noqa: BLE001
            logger.error("Image OCR extraction failed: %s", str(e))
            return {
                'text': '',
                'method': 'error',
                'success': False,
                'error': str(e)
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from PDF file with automatic OCR fallback.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extraction result dictionary
        """
        logger.info("Extracting text from PDF: %s", file_path.name)
        
        # Multimodal pipeline: PyMuPDF for layout + bounding boxes, pdfplumber for text, Camelot tables, OCR fallback when needed
        text = ''
        method = 'pdfplumber'
        layout_blocks = []
        tables = []
        pages = []

        # Step 0: capture layout with PyMuPDF (page-level text + blocks)
        if self.use_pymupdf:
            pages, pymupdf_text = self._extract_with_pymupdf(file_path)
        else:
            pymupdf_text = ''

        # Step 1: pdfplumber first (text-based PDFs)
        text = self._extract_with_pdfplumber(file_path)

        # Optional secondary extractor only if explicitly configured
        if not text.strip() and self.pdf_method == 'pypdf2':
            text = self._extract_with_pypdf2(file_path)
            method = 'pypdf2'

        # If pdfplumber returned nothing but PyMuPDF has text, use that text while keeping layout
        if not text.strip() and pymupdf_text.strip():
            text = pymupdf_text
            method = 'pymupdf'

        # Optional table extraction with Camelot (layout-aware)
        if self.use_camelot:
            tables = self._extract_tables_with_camelot(file_path)
        
        # Step 2: OCR fallback only when no text extracted
        if not text.strip() and self.ocr_fallback:
            logger.info("No text extracted; running OCR fallback")
            ocr_text = self._extract_with_ocr(file_path)
            if ocr_text.strip():
                logger.info("OCR produced %d chars", len(ocr_text.strip()))
                return {
                    'text': ocr_text,
                    'method': f'{method}_ocr_fallback',
                    'success': True,
                    'error': None,
                    'layout_blocks': layout_blocks,
                    'tables': tables,
                    'ocr_performed': True,
                    'pages': pages
                }

        success = len(text.strip()) > 0
        
        return {
            'text': text,
            'method': method,
            'success': success,
            'error': None if success else 'No text extracted',
            'layout_blocks': layout_blocks,
            'tables': tables,
            'ocr_performed': False,
            'pages': pages
        }

    def _extract_with_pymupdf(self, file_path: Path) -> tuple[list, str]:
        """Extract text with PyMuPDF; capture layout blocks and bounding boxes."""
        try:
            import fitz  # PyMuPDF
        except Exception as exc:  # PyMuPDF optional
            logger.debug("PyMuPDF unavailable: %s", exc)
            return [], ''

        pages = []
        text_parts = []

        try:
            with fitz.open(file_path) as doc:
                for page_index, page in enumerate(doc):
                    page_text = page.get_text("text") or ""
                    text_parts.append(page_text)

                    blocks = []
                    if self.capture_layout:
                        for blk in page.get_text("blocks"):
                            # blk: (x0, y0, x1, y1, text, block_no, block_type, ...)
                            if len(blk) >= 5:
                                blocks.append({
                                    'bbox': [blk[0], blk[1], blk[2], blk[3]],
                                    'text': blk[4]
                                })

                    pages.append({
                        'page_number': page_index + 1,
                        'text': page_text,
                        'blocks': blocks
                    })

            return pages, '\n'.join(text_parts)
        except Exception as exc:  # noqa: BLE001
            logger.warning("PyMuPDF extraction failed: %s", exc)
            return [], ''
    
    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """
        Extract text using pdfplumber.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            logger.warning("pdfplumber extraction failed: %s", str(e))
            return ''
    
    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """
        Extract text using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            text_parts = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return '\n'.join(text_parts)
        except Exception as e:
            logger.warning("PyPDF2 extraction failed: %s", str(e))
            return ''

    def _extract_tables_with_camelot(self, file_path: Path) -> list:
        """Extract tables using Camelot when available."""
        try:
            import camelot
        except Exception as exc:  # Camelot optional
            logger.debug("Camelot unavailable: %s", exc)
            return []

        try:
            tables = camelot.read_pdf(str(file_path), pages='1-end')
            return [tbl.df.to_dict(orient='records') for tbl in tables]
        except Exception as exc:  # noqa: BLE001
            logger.warning("Camelot extraction failed: %s", exc)
            return []
    
    def _extract_with_ocr(self, file_path: Path) -> str:
        """
        Extract text using OCR (pytesseract).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            from pdf2image import convert_from_path
            # Point pytesseract to bundled binary when configured
            if self.tesseract_path:
                pytesseract.pytesseract.tesseract_cmd = self.tesseract_path
            
            # Convert PDF pages to images
            images = convert_from_path(
                file_path,
                last_page=self.ocr_page_limit,
                poppler_path=self.poppler_path
            )
            
            text_parts = []
            for i, image in enumerate(images):
                logger.debug("Running OCR on page %d", i + 1)
                page_text = pytesseract.image_to_string(image, lang=self.ocr_language)
                if page_text:
                    text_parts.append(page_text)
            
            return '\n'.join([t for t in text_parts if len(t.strip()) >= self.ocr_min_chars])
            
        except ImportError:
            logger.error("pdf2image not installed. OCR fallback unavailable.")
            logger.info("Install with: pip install pdf2image")
            return ''
        except Exception as e:
            logger.error("OCR extraction failed: %s", str(e))
            return ''
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extraction result dictionary
        """
        logger.info("Extracting text from DOCX: %s", file_path.name)
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            text = '\n'.join(text_parts)
            success = len(text.strip()) >= self.min_text_length
            
            return {
                'text': text,
                'method': 'python-docx',
                'success': success,
                'error': None if success else 'Insufficient text extracted'
            }
            
        except Exception as e:
            logger.error("DOCX extraction failed: %s", str(e))
            return {
                'text': '',
                'method': 'error',
                'success': False,
                'error': str(e)
            }
    
    def extract_batch(self, file_paths: list) -> Dict[str, Dict]:
        """
        Extract text from multiple files.
        
        Args:
            file_paths: List of file paths
            
        Returns:
            Dictionary mapping file names to extraction results
        """
        logger.info("Starting batch extraction for %d files", len(file_paths))
        
        results = {}
        for file_path in file_paths:
            file_path = Path(file_path)
            results[file_path.name] = self.extract(file_path)
        
        success_count = sum(1 for r in results.values() if r['success'])
        logger.info("Batch extraction complete: %d/%d successful", 
                   success_count, len(file_paths))
        
        return results


def extract_text(file_path: Union[str, Path]) -> str:
    """
    Convenience function to extract text from a file.
    
    Args:
        file_path: Path to file
        
    Returns:
        Extracted text
    """
    extractor = TextExtractor()
    result = extractor.extract(file_path)
    return result['text'] if result['success'] else ''
